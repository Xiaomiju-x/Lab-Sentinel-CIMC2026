#!/usr/bin/env python3
"""Inventory downloaded CC-BY phosphor supplementary files without promoting labels.

This is a source-intake audit.  It extracts searchable text and structural facts
from PDF/DOCX files, but it deliberately does not digitize plots, infer labels, or
turn repeated curve points into independent experimental groups.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader


KEYWORDS = (
    "concentration",
    "quenching",
    "temperature",
    "thermal",
    "activation energy",
    "intensity",
    "lifetime",
    "quantum yield",
    "table",
    "raw data",
    "supporting data",
)


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest().upper()


def normalized_excerpt(text: str, start: int, width: int = 260) -> str:
    lo = max(0, start - width // 3)
    hi = min(len(text), start + width)
    return re.sub(r"\s+", " ", text[lo:hi]).strip()


def keyword_hits(text: str) -> list[dict[str, Any]]:
    hits: list[dict[str, Any]] = []
    lower = text.lower()
    for keyword in KEYWORDS:
        positions = [match.start() for match in re.finditer(re.escape(keyword), lower)]
        if positions:
            hits.append(
                {
                    "keyword": keyword,
                    "count": len(positions),
                    "first_excerpt": normalized_excerpt(text, positions[0]),
                }
            )
    return hits


def inspect_pdf(path: Path, root: Path) -> dict[str, Any]:
    reader = PdfReader(str(path))
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(
            {
                "page_number": page_number,
                "text_characters": len(text),
                "keyword_hits": keyword_hits(text),
            }
        )
    return {
        "relative_path": path.relative_to(root).as_posix(),
        "kind": "PDF",
        "bytes": path.stat().st_size,
        "sha256": sha256(path),
        "page_count": len(reader.pages),
        "pages": pages,
    }


def inspect_docx(path: Path, root: Path) -> dict[str, Any]:
    document = Document(str(path))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    tables = []
    for table_number, table in enumerate(document.tables, start=1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        tables.append(
            {
                "table_number": table_number,
                "row_count": len(rows),
                "column_count_max": max((len(row) for row in rows), default=0),
                "first_rows": rows[:5],
            }
        )
    full_text = "\n".join(paragraphs + [" | ".join(row) for t in tables for row in t["first_rows"]])
    return {
        "relative_path": path.relative_to(root).as_posix(),
        "kind": "DOCX",
        "bytes": path.stat().st_size,
        "sha256": sha256(path),
        "paragraph_count_nonempty": len(paragraphs),
        "table_count": len(tables),
        "tables": tables,
        "keyword_hits": keyword_hits(full_text),
        "paragraphs": paragraphs,
        "render_gate": "STRUCTURAL_ONLY_LIBREOFFICE_NOT_AVAILABLE",
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    root = args.root.resolve()
    records: list[dict[str, Any]] = []
    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            records.append(inspect_pdf(path, root))
        elif suffix == ".docx":
            records.append(inspect_docx(path, root))
        elif suffix == ".xls":
            records.append(
                {
                    "relative_path": path.relative_to(root).as_posix(),
                    "kind": "XLS",
                    "bytes": path.stat().st_size,
                    "sha256": sha256(path),
                    "inspection_state": "PENDING_ARTIFACT_TOOL_CONVERSION_AND_REVIEW",
                }
            )
    payload = {
        "schema": "pmc-phosphor-supplement-source-audit-v1",
        "claim_boundary": (
            "Inventory only. No plot digitization, teacher labels, publication text labels, "
            "or repeated curve-point pseudo-replication is permitted."
        ),
        "root": root.as_posix(),
        "file_count": len(records),
        "files": records,
    }
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(json.dumps({"output": str(args.output), "file_count": len(records)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
